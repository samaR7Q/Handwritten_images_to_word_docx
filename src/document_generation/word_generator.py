from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

class WordGenerator:
    """Generate Word documents from processed text"""
    
    def __init__(self):
        self.output_dir = "outputs"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def create_document(self, content, title="Converted Notes", diagrams=None):
        """
        Create Word document from structured content
        
        Args:
            content: Text content
            title: Document title
            diagrams: List of diagram dicts with 'path', 'bbox', 'center_y'
        """
        doc = Document()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # If we have diagrams, we need to insert them at appropriate positions
        if diagrams:
            self._create_document_with_diagrams(doc, content, diagrams)
        else:
            self._add_text_content(doc, content)
        
        # Save document
        output_path = os.path.join(self.output_dir, f"{title.replace(' ', '_')}.docx")
        doc.save(output_path)
        
        print(f"✅ Document saved: {output_path}")
        return output_path
    
    def _add_text_content(self, doc, content):
        """Add text content to document"""
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check for headings
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            
            # Check for diagrams
            elif '[DIAGRAM]' in line:
                p = doc.add_paragraph()
                p.add_run('[Diagram placeholder - original image preserved]').italic = True
            
            # Check for equations/formulas (basic detection)
            elif self._contains_formula(line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            
            # Regular text
            else:
                doc.add_paragraph(line)
    
    def _create_document_with_diagrams(self, doc, content, diagrams):
        """Create document with diagrams embedded at appropriate positions"""
        # Add text content first
        self._add_text_content(doc, content)
        
        # Add diagrams section
        doc.add_page_break()
        doc.add_heading("Diagrams and Figures", level=1)
        
        for idx, diagram in enumerate(diagrams):
            # Add diagram heading
            doc.add_heading(f"Figure {idx + 1}", level=2)
            
            # Add diagram image
            if os.path.exists(diagram['path']):
                try:
                    # Calculate width (max 6 inches)
                    doc.add_picture(diagram['path'], width=Inches(6))
                    
                    # Add caption
                    caption = doc.add_paragraph()
                    caption_run = caption.add_run(f"Figure {idx + 1}: Extracted diagram")
                    caption_run.italic = True
                    caption_run.font.size = Pt(10)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add spacing
                    doc.add_paragraph()
                    
                except Exception as e:
                    doc.add_paragraph(f"[Could not embed diagram: {e}]")
            else:
                doc.add_paragraph(f"[Diagram file not found: {diagram['path']}]")
    
    def _contains_formula(self, text):
        """Detect if text contains chemical/math formulas"""
        formula_patterns = [
            r'\d+[A-Z][a-z]?\d*',  # Chemical formulas like H2O, CaCl2
            r'[ΔΣ∫∂]',  # Math symbols
            r'[A-Za-z]+\d+',  # Subscripts
            r'→|⇌|≈|≠|≤|≥',  # Arrows and operators
        ]
        
        return any(re.search(pattern, text) for pattern in formula_patterns)
    
    def embed_image(self, doc, image_path, width=Inches(4)):
        """Add image to document"""
        try:
            doc.add_picture(image_path, width=width)
            return True
        except Exception as e:
            print(f"⚠️ Could not embed image: {e}")
            return False